"""
Document text extraction for resume parsing.
Handles PDF and DOCX file parsing with text cleaning and normalization.
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import logging

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from app.core.logging import get_logger
from app.utils.file_utils import get_file_info

logger = get_logger("text_extractor")


class TextExtractor:
    """Extract and clean text from resume documents."""
    
    def __init__(self):
        """Initialize the text extractor."""
        self.supported_formats = []
        
        if PDF_AVAILABLE:
            self.supported_formats.append('.pdf')
        if DOCX_AVAILABLE:
            self.supported_formats.append('.docx')
        
        # Always support plain text
        self.supported_formats.append('.txt')
        
        logger.info(f"TextExtractor initialized with support for: {self.supported_formats}")
    
    def extract_text(self, file_path: Path) -> Dict[str, any]:
        """
        Extract text from a resume file.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_info = get_file_info(file_path)
        file_ext = file_path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        logger.info(f"Extracting text from {file_path.name}")
        
        try:
            if file_ext == '.pdf':
                text_data = self._extract_pdf_text(file_path)
            elif file_ext == '.docx':
                text_data = self._extract_docx_text(file_path)
            elif file_ext == '.txt':
                text_data = self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Clean and normalize the extracted text
            cleaned_text = self._clean_text(text_data['raw_text'])
            
            result = {
                'file_path': str(file_path),
                'file_info': file_info,
                'raw_text': text_data['raw_text'],
                'cleaned_text': cleaned_text,
                'extraction_method': text_data['method'],
                'extraction_confidence': text_data.get('confidence', 1.0),
                'sections': text_data.get('sections', {}),
                'layout_info': text_data.get('layout', {})
            }
            
            logger.info(f"Successfully extracted {len(cleaned_text)} characters from {file_path.name}")
            return result
            
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path.name}: {e}")
            raise
    
    def _extract_pdf_text(self, file_path: Path) -> Dict[str, any]:
        """Extract text from PDF file using multiple methods."""
        if not PDF_AVAILABLE:
            raise ImportError("PDF processing libraries not available")
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                pages = []
                sections = {}
                
                for page_num, page in enumerate(pdf.pages):
                    # Extract text with layout preservation
                    text = page.extract_text()
                    if text:
                        pages.append(text)
                    
                    # Extract tables if present
                    tables = page.extract_tables()
                    if tables:
                        sections[f'tables_page_{page_num}'] = tables
                
                raw_text = '\n\n'.join(pages)
                
                return {
                    'raw_text': raw_text,
                    'method': 'pdfplumber',
                    'confidence': 0.9,
                    'sections': sections,
                    'layout': {
                        'page_count': len(pdf.pages),
                        'has_tables': bool(sections)
                    }
                }
                
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    pages = []
                    
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            pages.append(text)
                    
                    raw_text = '\n\n'.join(pages)
                    
                    return {
                        'raw_text': raw_text,
                        'method': 'PyPDF2',
                        'confidence': 0.7,
                        'layout': {
                            'page_count': len(pdf_reader.pages)
                        }
                    }
                    
            except Exception as e2:
                logger.error(f"Both PDF extraction methods failed: {e2}")
                raise
    
    def _extract_docx_text(self, file_path: Path) -> Dict[str, any]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX processing library not available")
        
        try:
            doc = Document(file_path)
            paragraphs = []
            sections = {}
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            if tables:
                sections['tables'] = tables
            
            raw_text = '\n\n'.join(paragraphs)
            
            return {
                'raw_text': raw_text,
                'method': 'python-docx',
                'confidence': 0.95,
                'sections': sections,
                'layout': {
                    'paragraph_count': len(paragraphs),
                    'table_count': len(tables)
                }
            }
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    def _extract_txt_text(self, file_path: Path) -> Dict[str, any]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                raw_text = file.read()
            
            return {
                'raw_text': raw_text,
                'method': 'plain_text',
                'confidence': 1.0,
                'layout': {
                    'line_count': len(raw_text.split('\n'))
                }
            }
            
        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned and normalized text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with NLP
        text = re.sub(r'[^\w\s\-\.\,\;\:\!\?\(\)\[\]\{\}\@\#\$\%\&\*\+\=\/\|\\]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def identify_sections(self, text: str) -> Dict[str, str]:
        """
        Identify resume sections based on common headers.
        
        Args:
            text: Cleaned resume text
            
        Returns:
            Dictionary mapping section names to section content
        """
        sections = {}
        
        # Common resume section headers
        section_patterns = {
            'contact': r'(?i)(contact|personal|info|information)',
            'summary': r'(?i)(summary|profile|objective|about)',
            'experience': r'(?i)(experience|work\s+history|employment|career)',
            'education': r'(?i)(education|academic|qualifications)',
            'skills': r'(?i)(skills|competencies|technologies|tools)',
            'projects': r'(?i)(projects|portfolio|achievements)',
            'certifications': r'(?i)(certifications|certificates|licenses)',
            'languages': r'(?i)(languages|language\s+skills)',
            'interests': r'(?i)(interests|hobbies|activities)'
        }
        
        lines = text.split('\n')
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            section_found = False
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line, re.IGNORECASE):
                    # Save previous section
                    if current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    
                    # Start new section
                    current_section = section_name
                    current_content = []
                    section_found = True
                    break
            
            if not section_found:
                current_content.append(line)
        
        # Save the last section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    def get_extraction_stats(self, text_data: Dict[str, any]) -> Dict[str, any]:
        """
        Get statistics about the extracted text.
        
        Args:
            text_data: Result from extract_text method
            
        Returns:
            Dictionary with extraction statistics
        """
        raw_text = text_data.get('raw_text', '')
        cleaned_text = text_data.get('cleaned_text', '')
        
        stats = {
            'raw_length': len(raw_text),
            'cleaned_length': len(cleaned_text),
            'word_count': len(cleaned_text.split()),
            'line_count': len(cleaned_text.split('\n')),
            'extraction_method': text_data.get('extraction_method', 'unknown'),
            'confidence': text_data.get('extraction_confidence', 0.0),
            'has_tables': bool(text_data.get('sections', {}).get('tables')),
            'section_count': len(text_data.get('sections', {}))
        }
        
        return stats
